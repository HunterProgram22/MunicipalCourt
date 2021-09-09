import win32com.client
import datetime
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH


class calendarEvent(object):
    def __init__(self, start_time, subject, body):
        start_time = str(start_time)
        # print(start_time)
        self.start_time = start_time[:16]
        # print(self.start_time)
        self.subject = subject
        self.body = body
        self.sanitize_body()

    def sanitize_body(self):
        if "-~" in self.body:
            new_body = self.body.split("-~")
            new_body = new_body[0].rstrip()
        else:
            new_body = self.body.split("--")
            new_body = new_body[0].rstrip()
        self.body = new_body


def getCalendarEntries(room, days=1):
    """
    Returns calender entries for days default is 1
    """
    Outlook = win32com.client.Dispatch("Outlook.Application")
    ns = Outlook.GetNamespace("MAPI")
    # turn this into a list to read more calendars
    recipient = ns.CreateRecipient(room)  # cmd whoami to find this
    resolved = recipient.Resolve()  # checks for username in address book
    # olFolderCalendar = 9
    # appointments = ns.GetDefaultFolder(9).Items  # for personal calendar
    appointments = ns.GetSharedDefaultFolder(recipient, 9).Items
    # appointments = ns.GetDefaultFolder(9).Items
    print(appointments)
    appointments.Sort("[Start]")
    appointments.IncludeRecurrences = "True"
    today = datetime.datetime.today()
    begin = today.date().strftime("%m/%d/%Y")
    tomorrow = datetime.timedelta(days=days) + today
    end = tomorrow.date().strftime("%m/%d/%Y")
    appointments = appointments.Restrict(
        "[Start] >= '" + begin + "' AND [END] <= '" + end + "'"
    )
    # events={'Start':[],'Subject':[],'Duration':[], 'Body':[]}
    events = []
    for a in appointments:
        new_event = calendarEvent(a.Start, a.Subject, a.Body)
        events.append(new_event)
    return events


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def create_daily_calendar(courtroom):
    events = getCalendarEntries(courtroom)
    mydoc = docx.Document()
    court_names = {
        "courtrooma": "Courtroom A",
        "courtroomb": "Courtoom B",
        "courtroomc": "Courtroom C",
    }
    if len(events) == 0:
        heading = mydoc.add_heading(
            court_names[courtroom]
            + " "
            + str(datetime.datetime.now().strftime("%m-%d-%Y"))
        )
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph = mydoc.add_paragraph("Nothing scheduled today.")
    else:
        heading = mydoc.add_heading(
            court_names[courtroom] + " " + str(events[0].start_time[:10] + "\n")
        )
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for event in events:
            paragraph = mydoc.add_paragraph()
            time = paragraph.add_run(event.start_time[10:] + " \n")
            time.bold = True
            paragraph.add_run(event.subject + " " + event.body + "\n")
    # mydoc.save("C:\\users\\jkudela\\Desktop\\" + courtroom + ".docx")
    today = datetime.datetime.today()
    begin = today.date().strftime("%m-%d-%Y")
    mydoc.save("J:\\Courtroom_Calendars\\" + courtroom + "_" + begin + ".docx")
    # mydoc.save("J:\\Courtroom_Calendars\\" + courtroom + "_" + begin + "test.docx")


create_daily_calendar("courtrooma")
create_daily_calendar("courtroomb")
create_daily_calendar("courtroomc")
