import win32com.client
import datetime
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH


class calendarEvent(object):
    def __init__(self, start_time, subject, body):
        """The start time of the calendar event is converted to a string, then
        it is stripped of all extraneous time information (seconds, timezone,
        etc.). The body is 'sanitized' by removing all the webex information
        using symbols that appear right before the webex information starts."""
        start_time = str(start_time)
        self.start_time = start_time[:16]
        self.subject = subject
        self.body = body
        self.sanitize_body()

    def sanitize_body(self):
        if "-~" in self.body:
            new_body = self.body.split("-~")
            new_body = new_body[0].rstrip()
        else:
            new_body = self.body.split("--")
            new_body = new_body[0].rstrip()
        self.body = new_body


def getCalendarEntries(room, days=1):
    """
    Returns calender entries for days default is 1
    This function is modified to get tomorrow's calendar. An extra day is
    added to today so start time is tomorrow.
    """
    Outlook = win32com.client.Dispatch("Outlook.Application")
    ns = Outlook.GetNamespace("MAPI")
    recipient = ns.CreateRecipient(room)  # cmd whoami to find this
    resolved = recipient.Resolve()  # checks for username in address book
    appointments = ns.GetSharedDefaultFolder(recipient, 9).Items
    print(appointments)
    appointments.Sort("[Start]")
    appointments.IncludeRecurrences = "True"
    today = datetime.datetime.today() + datetime.timedelta(days)
    begin = today.date().strftime("%m/%d/%Y")
    tomorrow = datetime.timedelta(days) + today
    end = tomorrow.date().strftime("%m/%d/%Y")
    appointments = appointments.Restrict(
        "[Start] >= '" + begin + "' AND [END] <= '" + end + "'"
    )
    events = []
    for a in appointments:
        new_event = calendarEvent(a.Start, a.Subject, a.Body)
        events.append(new_event)
    return events


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def create_daily_calendar(courtroom):
    events = getCalendarEntries(courtroom)
    mydoc = docx.Document()
    court_names = {
        "courtrooma": "Courtroom A",
        "courtroomb": "Courtoom B",
        "courtroomc": "Courtroom C",
    }
    if len(events) == 0:
        heading = mydoc.add_heading(
            court_names[courtroom]
            + " "
            + str(datetime.datetime.now().strftime("%m-%d-%Y"))
        )
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph = mydoc.add_paragraph("Nothing scheduled today.")
    else:
        heading = mydoc.add_heading(
            court_names[courtroom] + " " + str(events[0].start_time[:10] + "\n")
        )
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for event in events:
            paragraph = mydoc.add_paragraph()
            time = paragraph.add_run(event.start_time[10:] + " \n")
            time.bold = True
            paragraph.add_run(event.subject + " " + event.body + "\n")
    today = datetime.datetime.today() + datetime.timedelta(1)
    begin = today.date().strftime("%m-%d-%Y")
    mydoc.save("J:\\Courtroom_Calendars\\" + courtroom + "_clerk_" + begin + ".docx")


create_daily_calendar("courtrooma")
create_daily_calendar("courtroomb")
create_daily_calendar("courtroomc")
