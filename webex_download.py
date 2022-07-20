import urllib
import docx
import datetime
import requests
import json
import os

from flatten_json import flatten
from loguru import logger

import docx2pdf
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

VIDEO_PATH = "V:\\COURTROOM VIDEO PROCEEDINGS - JULY 2021 TO PRESENT\\"
BEARER_TOKEN = "YzdkNjYyMzktZTBjMi00NWNjLThjMDgtNDFkNmQwZTFlNjVkYjI5MDEwNDgtODdk_PF84_6b8180ed-7e73-4208-90f5-b67a07de84ac"

def add_hyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(
        docx.oxml.shared.qn("r:id"),
        r_id,
    )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement("w:r")

    # Create a new w:rPr element
    rPr_element = docx.oxml.shared.OxmlElement("w:rPr")

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr_element)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def create_single_day_court_recordings(response_items_delaware, response_items_municipal, start_date, month, year):
    mydoc = docx.Document()
    start_date_list = start_date.split("-")
    heading_date = (
        start_date_list[1] + "-" + start_date_list[2] + "-" + start_date_list[0]
    )
    heading = mydoc.add_heading("Court Video Proceedings " + heading_date + "\n")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    month_dict = {'12': "December", '11': "November", '01': "January", '02': "February", '03': "March", '04': "April", '05': "May", '06':"June", '07':"July", '08': "August", '09': "September", '10': "October"}
    month = month_dict.get(month)

    mydoc_paragraph = mydoc.add_paragraph()
    for video_info in response_items_delaware:
        playback_url = video_info["playbackUrl"]
        title = video_info["topic"]
        print(title)
        print(playback_url)
        run = mydoc_paragraph.add_run(title + "\n")
        run.bold
        run.bold = True
        run_link = add_hyperlink(mydoc_paragraph, playback_url, "Video Link")
        mydoc_paragraph.add_run("\n\n")
    for video_info in response_items_municipal:
        playback_url = video_info["playbackUrl"]
        title = video_info["topic"]
        print(title)
        print(playback_url)
        run = mydoc_paragraph.add_run(title + "\n")
        run.bold
        run.bold = True
        run_link = add_hyperlink(mydoc_paragraph, playback_url, "Video Link")
        mydoc_paragraph.add_run("\n\n")

    document_name = "Courtroom_Proceedings_" + heading_date + ".docx"
    mydoc.save(VIDEO_PATH + month + " " + year + "\\" + document_name)
    docx2pdf.convert(VIDEO_PATH + month + " " + year + "\\" + document_name)
    os.remove(VIDEO_PATH + month + " " + year + "\\" + document_name)



@logger.catch
def main():
    month = input("Enter month of video proceedings in MM format (i.e. '09 for September'): ")
    day_date = input(str("Enter the first day date of the video proceedings in DD format: "))
    year = input(str("Enter year of video proceedings (i.e. '2021'): "))

    # start_date = input(
    #     str("Enter the date (YYYY-MM-DD) of the first day of recordings to download:")
    # )
    # # end_date = input(
    #     str(
    #         "Enter the date (YYYY-MM-DD) of the last day of recordings to download"
    #         + " - download does not include the last day:"
    #     )
    # )

    start_date = year + '-' + month + '-' + day_date
    next_day_date = int(day_date) + 1
    end_date = year + '-' + month + '-' + str(next_day_date)
    # end_date = '2022-06-01' # Need to use manual last date if last day of month is a 31st TODO - Fix  ,s
    URL_string_delaware = (
        "https://webexapis.com/v1/recordings?max=100&from="
        + "{from_date}&to={to_date}&siteUrl={site_url}&hostEmail={email}".format(
            max_records="100",  # There is something up with max_records in the format
            from_date=start_date,
            to_date=end_date,
            site_url="delawareohio.webex.com",
            email="jkudela@municipalcourt.org",
        )
    )

    URL_string_municipal = (
        "https://webexapis.com/v1/recordings?max=100&from="
        + "{from_date}&to={to_date}&siteUrl={site_url}&hostEmail={email}".format(
            max_records="100",  # There is something up with max_records in the format
            from_date=start_date,
            to_date=end_date,
            site_url="municipalcourt.webex.com",
            email="jkudela@municipalcourt.org",
        )
    )

    location = "Delaware City Webex"
    headers = {
        "Authorization": "Bearer {bearer_token}".format(bearer_token=BEARER_TOKEN)
    }
    response_delaware = requests.get(
        url=URL_string_delaware,
        headers=headers,
        stream=True,
    )

    response_municipal = requests.get(
        url=URL_string_municipal,
        headers=headers,
        stream=True,
    )

    response_data_delaware = response_delaware.json()
    response_data_municipal = response_municipal.json()

    response_items_delaware = response_data_delaware["items"]
    response_items_municipal = response_data_municipal["items"]

    create_single_day_court_recordings(response_items_delaware, response_items_municipal, start_date, month, year)


main()
